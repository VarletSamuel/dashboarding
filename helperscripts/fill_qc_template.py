"""Fill the Quality Controls template with managed-service scan results.

Workflow:
1. Optionally download manifest + report files from Azure Blob Storage
   (--storage-connection-string). If omitted, local files are used as-is.
2. Resolve the latest export folder and read manifest.json.
3. Load relevant CSV files referenced by manifest entries.
4. Replace scalar/table/recommendation anchors in the QC template.

Usage:
    pip install python-docx azure-storage-blob
    python fill_qc_template.py -c CUST --storage-connection-string "<conn>"
    python fill_qc_template.py -c CUST --storage-connection-string "<conn>" --template-version 1.1
    python fill_qc_template.py -c CUST                       # use local files only
    python fill_qc_template.py -c CUST --run-wrapper         # run extractors first (legacy)

Conventions used in the template:
    {{NAME}}                    - simple scalar; appears alone on a line
                                  OR in front-matter table cells
    {{TABLE_<RESOURCE>}}        - insert a findings table at this anchor
    {{RECOMMENDATIONS_<RES>}}   - replace with a recommendation paragraph
                                  (multi-line strings get split per line)
"""

from __future__ import annotations
import argparse
import csv
import json
import os
from pathlib import Path
import subprocess
import sys
from copy import deepcopy
from datetime import date, datetime, timezone
from urllib.parse import urlparse
from docx import Document
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT_DIR = SCRIPT_DIR.parent
DEFAULT_TEMPLATE = ROOT_DIR / "templates" / "YYYY-MM-DD_CUSTOMER_QualityControls.docx"
DEFAULT_TEMPLATE_DIR = ROOT_DIR / "templates"
DEFAULT_REPORTS_DIR = ROOT_DIR / "reports"

# Table style available in the delaware template. Confirmed by inspecting the
# source CostManagement.docx — keep this in sync if you rename the style.
DLW_TABLE_STYLE = "DLWTableTeal"


# -----------------------------------------------------------------------------
# Replacement helpers
# -----------------------------------------------------------------------------

def _iter_paragraphs(doc):
    """Yield every paragraph in the document, including those nested in
    tables AND those wrapped in <w:sdt> structured document tags (the cover
    page title/subtitle/date use SDTs). python-docx's `doc.paragraphs` only
    returns direct <w:p> children of <w:body>, so we walk the XML ourselves
    to catch SDT-wrapped paragraphs."""
    from docx.text.paragraph import Paragraph
    # Body paragraphs, including those inside <w:sdt><w:sdtContent>...
    body = doc.element.body
    for p_el in body.iter(qn("w:p")):
        yield Paragraph(p_el, doc.element.body)
    # Note: the above also yields paragraphs that live inside tables, since
    # lxml's iter() walks the whole subtree. No need for a separate table pass.


def _iter_anchor_keys(doc, prefix: str) -> list[str]:
    """Return unique placeholder keys for paragraphs like {{PREFIX_*}}."""
    keys: list[str] = []
    seen: set[str] = set()
    marker = "{{" + prefix + "_"

    for p in _iter_paragraphs(doc):
        text = p.text.strip()
        if not (text.startswith(marker) and text.endswith("}}")):
            continue
        key = text[2:-2]
        if key not in seen:
            seen.add(key)
            keys.append(key)
    return keys


def replace_scalar(doc, key: str, value: str) -> int:
    """Replace `{{KEY}}` with `value` everywhere it appears.

    Walks every `<w:t>` element in the document body so it also catches text
    inside SDT-wrapped runs (cover-page title/subtitle/date). Cross-run
    placeholders are handled by collapsing all `<w:t>` text in a paragraph
    when no single `<w:t>` contains the needle on its own."""
    needle = f"{{{{{key}}}}}"
    count = 0
    body = doc.element.body
    for p_el in body.iter(qn("w:p")):
        # Fast path: find any <w:t> whose text contains the full needle.
        t_elements = list(p_el.iter(qn("w:t")))
        joined = "".join(t.text or "" for t in t_elements)
        if needle not in joined:
            continue
        # Try a single-<w:t> hit first (preserves formatting of other runs).
        hit = False
        for t in t_elements:
            if t.text and needle in t.text:
                t.text = t.text.replace(needle, value)
                hit = True
                count += 1
                break
        if hit:
            continue
        # Cross-<w:t> placeholder: write the full collapsed text into the
        # first <w:t> and blank the rest. Acceptable because our placeholders
        # always sit in single-format spans.
        t_elements[0].text = joined.replace(needle, value)
        for t in t_elements[1:]:
            t.text = ""
        count += 1
    return count


def _find_anchor_paragraph(doc, key: str):
    """Return the paragraph element whose text equals `{{KEY}}` (exact match,
    stripped). The template puts each table/recommendation placeholder on its
    own paragraph, so an exact match is what we want."""
    needle = f"{{{{{key}}}}}"
    for p in _iter_paragraphs(doc):
        if p.text.strip() == needle:
            return p
    return None


def insert_table_at(doc, key: str, headers: list[str], rows: list[list]) -> bool:
    """Replace the `{{KEY}}` anchor paragraph with a real Word table.

    Returns False when the anchor isn't found (so the script can warn loudly
    rather than silently dropping data)."""
    anchor = _find_anchor_paragraph(doc, key)
    if anchor is None:
        return False

    # Build the table somewhere safe, then move its XML before the anchor.
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tbl.style = DLW_TABLE_STYLE
    except KeyError:
        # Style missing for some reason — fall back to default grid.
        tbl.style = "Table Grid"

    # Header row
    for col_idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[col_idx]
        cell.text = str(header)
        # First-row conditional formatting is driven by the table style; the
        # DLW table styles already paint the header band, no extra work here.

    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            tbl.rows[r].cells[c].text = "" if value is None else str(value)

    # Move the freshly-added table to sit right before the anchor paragraph,
    # then delete the anchor itself.
    tbl_el = tbl._element
    anchor._element.addprevious(tbl_el)
    anchor._element.getparent().remove(anchor._element)
    return True


def replace_recommendation(doc, key: str, text: str) -> bool:
    """Replace `{{RECOMMENDATIONS_*}}` with one or more paragraphs of text."""
    anchor = _find_anchor_paragraph(doc, key)
    if anchor is None:
        return False

    lines = text.splitlines() or [""]
    # Reuse the anchor paragraph for the first line so its style is preserved.
    anchor.runs[0].text = lines[0]
    for extra in lines[1:]:
        new_p = deepcopy(anchor._element)
        # Strip existing runs in the copy and add a clean one.
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        anchor._element.addnext(new_p)
        # Re-wrap the copy as a Paragraph object isn't directly exposed — easier:
        # walk python-docx's paragraph list again to find the new one.
        # Simpler approach: just append a run via raw XML.
        run = anchor._element.makeelement(qn("w:r"), {})
        t = anchor._element.makeelement(qn("w:t"), {})
        t.text = extra
        run.append(t)
        new_p.append(run)
    return True


# -----------------------------------------------------------------------------
# Data sourcing from managed-service exports
# -----------------------------------------------------------------------------

def _normalize_token(value: str) -> str:
    return "".join(ch for ch in value.upper() if ch.isalnum())


def _read_json(path: Path) -> dict:
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def _parse_timestamp(value: str) -> datetime:
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return datetime.min.replace(tzinfo=timezone.utc)


def _read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with open(path, "r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle, delimiter=";")
        return [dict(row) for row in reader]


def _find_latest_export_dir(customer: str, reports_dir: Path) -> Path:
    candidates: list[tuple[datetime, Path]] = []

    direct = reports_dir / customer
    if direct.exists() and direct.is_dir() and (direct / "manifest.json").exists():
        manifest = _read_json(direct / "manifest.json")
        stamp = _parse_timestamp(str(manifest.get("generated_at", "")))
        candidates.append((stamp, direct))

    prefix = f"{customer}_"
    for child in reports_dir.iterdir():
        if not child.is_dir() or not child.name.startswith(prefix):
            continue
        manifest_path = child / "manifest.json"
        if not manifest_path.exists():
            continue
        manifest = _read_json(manifest_path)
        stamp = _parse_timestamp(str(manifest.get("generated_at", "")))
        candidates.append((stamp, child))

    if not candidates:
        raise SystemExit(f"No export folder with manifest.json found for customer {customer} in {reports_dir}")

    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _is_sas_url(value: str) -> bool:
    parsed = urlparse(value)
    if parsed.scheme not in {"http", "https"}:
        return False
    return "sig=" in (parsed.query or "").lower()


def _download_reports_from_blob(
    customer: str,
    reports_dir: Path,
    connection_input: str,
    container_name: str,
    blob_prefix: str,
) -> None:
    """Download manifest.json and all files it references from blob storage.

    Files are written into *reports_dir / customer* so that
    ``_find_latest_export_dir`` can locate them normally afterwards.
    Already-present local files are skipped (not re-downloaded).
    """
    try:
        from azure.storage.blob import BlobServiceClient  # type: ignore[import-not-found]
    except ImportError as exc:
        raise SystemExit(
            "Missing Azure Blob SDK. Install with: pip install azure-storage-blob"
        ) from exc

    if _is_sas_url(connection_input):
        parsed = urlparse(connection_input)
        account_url = f"{parsed.scheme}://{parsed.netloc}"
        sas_token = parsed.query
        blob_service = BlobServiceClient(account_url=account_url, credential=sas_token)
    else:
        blob_service = BlobServiceClient.from_connection_string(connection_input)

    prefix = blob_prefix.strip("/")
    export_dir = reports_dir / customer
    export_dir.mkdir(parents=True, exist_ok=True)

    # --- Download manifest.json first so we know what else to fetch ----------
    manifest_blob = f"{prefix}/manifest.json" if prefix else "manifest.json"
    print(f"  Downloading manifest from blob: {manifest_blob}")
    try:
        manifest_bytes = blob_service.get_blob_client(
            container=container_name, blob=manifest_blob
        ).download_blob().readall()
    except Exception as exc:
        raise SystemExit(f"Could not download manifest from blob ({manifest_blob}): {exc}") from exc

    (export_dir / "manifest.json").write_bytes(manifest_bytes)
    manifest = json.loads(manifest_bytes.decode("utf-8"))

    # --- Collect referenced filenames ----------------------------------------
    filenames: list[str] = []
    for dashboard in manifest.get("dashboards", []):
        for file_spec in dashboard.get("files", []):
            fn = file_spec.get("filename")
            if fn:
                filenames.append(str(fn))
    for other in manifest.get("other_files", []):
        fn = other.get("filename")
        if fn:
            filenames.append(str(fn))

    downloaded = skipped = failed = 0
    for filename in filenames:
        local_path = export_dir / filename
        if local_path.exists():
            skipped += 1
            continue
        blob_name = f"{prefix}/{filename}" if prefix else filename
        try:
            data = blob_service.get_blob_client(
                container=container_name, blob=blob_name
            ).download_blob().readall()
            local_path.parent.mkdir(parents=True, exist_ok=True)
            local_path.write_bytes(data)
            downloaded += 1
        except Exception as exc:
            print(f"  Warning: could not download {blob_name}: {exc}")
            failed += 1

    print(
        f"  Blob download complete: {downloaded} downloaded, "
        f"{skipped} already present, {failed} failed."
    )


def _run_wrapper(
    customer: str,
    reports_dir: Path,
    skip_login: bool,
    date_from: str | None,
    date_to: str | None,
    lookback: str | None,
    sp_client_id: str | None,
    sp_client_secret: str | None,
    sp_certificate: str | None,
) -> None:
    wrapper_path = SCRIPT_DIR / "managedServiceWrapper.py"
    if not wrapper_path.exists():
        raise SystemExit(f"managedServiceWrapper.py not found at {wrapper_path}")

    cmd = [
        sys.executable,
        str(wrapper_path),
        "-c",
        customer,
        "--output-dir",
        str(reports_dir),
        "--output-format",
        "csv",
    ]
    if skip_login:
        cmd.append("--skip-login")
    if date_from:
        cmd.extend(["--from", date_from])
    if date_to:
        cmd.extend(["--to", date_to])
    if lookback and not date_from and not date_to:
        cmd.extend(["--lookback", lookback])
    if sp_client_id:
        cmd.extend(["--sp-client-id", sp_client_id])
    if sp_client_secret:
        cmd.extend(["--sp-client-secret", sp_client_secret])
    if sp_certificate:
        cmd.extend(["--sp-certificate", sp_certificate])

    print("Running managed service exports...")
    result = subprocess.run(cmd, cwd=str(ROOT_DIR), text=True)
    if result.returncode != 0:
        raise SystemExit(f"managedServiceWrapper failed with exit code {result.returncode}")


def _manifest_summary_file(manifest: dict, dashboard_id: str) -> str | None:
    for dashboard in manifest.get("dashboards", []):
        if dashboard.get("id") != dashboard_id:
            continue
        for file_spec in dashboard.get("files", []):
            if file_spec.get("type") == "summary":
                return file_spec.get("filename")
    return None


def _manifest_other_file(manifest: dict, file_type: str) -> str | None:
    for item in manifest.get("other_files", []):
        if item.get("type") == file_type:
            return item.get("filename")
    return None


def _find_latest_csv_by_fragment(export_dir: Path, fragment: str) -> Path | None:
    """Fallback for exports not yet represented in manifest (e.g. timestamp-named files)."""
    candidates = sorted(
        export_dir.glob(f"*{fragment}*.csv"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return candidates[0] if candidates else None


def _safe_int(value: str | None) -> int:
    if value is None:
        return 0
    try:
        return int(float(str(value).replace(",", ".").strip()))
    except (TypeError, ValueError):
        return 0


def _extract_template_version(path: Path) -> str | None:
    name = path.name
    marker = "_QualityControls_v"
    if marker not in name or not name.lower().endswith(".docx"):
        return None
    suffix = name.split(marker, 1)[1]
    return suffix[:-5] if suffix.lower().endswith(".docx") else suffix


def _version_sort_key(version: str) -> tuple[int, ...] | None:
    parts = version.split(".")
    if not parts or any(not p.isdigit() for p in parts):
        return None
    return tuple(int(p) for p in parts)


def _resolve_template_path(template_arg: str | None, template_version: str | None) -> Path:
    if template_arg:
        return Path(template_arg)

    if not DEFAULT_TEMPLATE_DIR.exists():
        raise SystemExit(f"Template directory not found: {DEFAULT_TEMPLATE_DIR}")

    candidates = sorted(DEFAULT_TEMPLATE_DIR.glob("*_QualityControls_v*.docx"))
    if not candidates:
        if DEFAULT_TEMPLATE.exists():
            return DEFAULT_TEMPLATE
        raise SystemExit(
            "No quality-control templates found under templates/. "
            "Expected files like *_QualityControls_v1.1.docx"
        )

    if template_version:
        wanted = template_version.strip().lower()
        for candidate in candidates:
            candidate_version = (_extract_template_version(candidate) or "").lower()
            if candidate_version == wanted:
                return candidate
        available = ", ".join(sorted(v for v in (_extract_template_version(p) for p in candidates) if v))
        raise SystemExit(
            f"Template version '{template_version}' not found. "
            f"Available versions: {available}"
        )

    numeric_candidates: list[tuple[tuple[int, ...], Path]] = []
    for candidate in candidates:
        version = _extract_template_version(candidate)
        if not version:
            continue
        key = _version_sort_key(version)
        if key is not None:
            numeric_candidates.append((key, candidate))

    if numeric_candidates:
        numeric_candidates.sort(key=lambda item: item[0], reverse=True)
        return numeric_candidates[0][1]

    raise SystemExit(
        "No numeric template versions found in templates/. "
        "Provide --template-version explicitly (for example: --template-version SYNH)."
    )


def _format_asp_sku_instances(row: dict[str, str]) -> str:
    """Build a human-readable SKU + instance summary for App Service Plans."""
    sku = (row.get("sku_name") or row.get("sku_size") or "n/a").strip() or "n/a"
    observed_instances = (row.get("observed_instance_peak_window") or "").strip()
    configured_instances = (row.get("sku_capacity") or "").strip()
    instances = observed_instances or configured_instances or "n/a"
    return f"{sku} ({instances} instance(s))"


def _build_law_recommendation(law_rows: list[dict[str, str]]) -> str:
    """Build a recommendation string from Log Analytics workspace summary rows."""
    total = len(law_rows)
    if total == 0:
        return (
            "No Log Analytics workspace data found.\n"
            "Ensure get_loganalyticsworkspace.py has been run and its output is present in the export folder."
        )
    cap_exceeded = [r for r in law_rows if (r.get("daily_quota_exceeded") or "").strip().lower() == "true"]
    no_cap = [r for r in law_rows if (r.get("daily_quota_gb") or "").strip() in {"N/A", "-1", ""}]
    unlimited = [r for r in law_rows if (r.get("daily_quota_gb") or "").strip().lower() == "unlimited"]
    lines = [
        f"Log Analytics workspaces scanned: {total}.",
        f"Workspaces with no daily cap configured: {len(no_cap)} — consider setting a cap to guard against runaway ingestion costs.",
        f"Workspaces with unlimited cap: {len(unlimited)} — validate commitment tier pricing is in place.",
        f"Workspaces where daily cap was exceeded: {len(cap_exceeded)} — review ingestion sources and consider raising the cap or filtering noisy tables.",
        "General guidance:",
        "- Align retention policy with compliance requirements; the default 30-day free tier may be insufficient.",
        "- Use commitment tiers (100 GB/day+) where average daily ingestion justifies it to reduce per-GB cost.",
        "- Enable workspace-level data export for long-term archival instead of extending retention inside Log Analytics.",
    ]
    return "\n".join(lines)


def _build_payload_from_exports(customer: str, export_dir: Path, manifest: dict) -> dict:
    subscriptions_file = _manifest_other_file(manifest, "subscriptions")
    subscriptions_rows: list[dict[str, str]] = []
    if subscriptions_file:
        subscriptions_rows = _read_csv(export_dir / subscriptions_file)

    appservice_file = _manifest_summary_file(manifest, "appServicePlans")
    appservice_rows = _read_csv(export_dir / appservice_file) if appservice_file else []

    keyvault_file = _manifest_summary_file(manifest, "keyVaults")
    keyvault_path = (export_dir / keyvault_file) if keyvault_file else None
    if keyvault_path is None or not keyvault_path.exists():
        keyvault_path = _find_latest_csv_by_fragment(export_dir, "keyvaults_summary")
    keyvault_rows = _read_csv(keyvault_path) if keyvault_path else []

    storage_file = _manifest_summary_file(manifest, "storageAccounts")
    storage_path = (export_dir / storage_file) if storage_file else None
    if storage_path is None or not storage_path.exists():
        storage_path = _find_latest_csv_by_fragment(export_dir, "storage_accounts_summary")
    storage_rows = _read_csv(storage_path) if storage_path else []

    sql_file = _manifest_summary_file(manifest, "sqlDatabases")
    sql_path = (export_dir / sql_file) if sql_file else None
    if sql_path is None or not sql_path.exists():
        sql_path = _find_latest_csv_by_fragment(export_dir, "sql_summary")
    sql_rows = _read_csv(sql_path) if sql_path else []

    quota_file = _manifest_summary_file(manifest, "quotaConsumption")
    quota_path = (export_dir / quota_file) if quota_file else None
    if quota_path is None or not quota_path.exists():
        quota_path = _find_latest_csv_by_fragment(export_dir, "quota_consumption_summary")
    quota_rows = _read_csv(quota_path) if quota_path else []

    entra_file = _manifest_summary_file(manifest, "appSecretExpirations")
    entra_path = (export_dir / entra_file) if entra_file else None
    if entra_path is None or not entra_path.exists():
        entra_path = _find_latest_csv_by_fragment(export_dir, "app_secret_expirations_summary")
    entra_rows = _read_csv(entra_path) if entra_path else []

    law_file = _manifest_summary_file(manifest, "logAnalyticsWorkspaces")
    law_path = (export_dir / law_file) if law_file else None
    if law_path is None or not law_path.exists():
        law_path = _find_latest_csv_by_fragment(export_dir, "loganalyticsworkspace_summary")
    law_rows = _read_csv(law_path) if law_path else []

    in_scope = [r for r in subscriptions_rows if (r.get("subscription_status") or "").lower() == "enabled"]
    out_scope = [r for r in subscriptions_rows if (r.get("subscription_status") or "").lower() != "enabled"]

    tenant_values = sorted({(r.get("tenant_id") or "").strip() for r in subscriptions_rows if r.get("tenant_id")})
    tenant_scalar = tenant_values[0] if len(tenant_values) == 1 else ", ".join(tenant_values) or "n/a"

    appservice_reco = [
        row for row in appservice_rows
        if (row.get("cost_signal_category") or "").strip() in {"orphaned", "stopped", "oversized-premium", "scale-in-candidate", "scale-up-review"}
    ]
    keyvault_low_quality = [
        row for row in keyvault_rows
        if _safe_float(row.get("quality_score")) < 60.0
    ]
    storage_low_quality = [
        row for row in storage_rows
        if _safe_float(row.get("quality_score")) < 60.0
    ]
    sql_pressure = [
        row for row in sql_rows
        if (row.get("advisory_category") or "").strip() in {"storage-pressure", "scale-up-review", "deadlock-review"}
    ]

    quota_by_subscription: dict[tuple[str, str], dict[str, object]] = {}
    for row in quota_rows:
        key = (
            (row.get("subscription_name") or "").strip(),
            (row.get("subscription_id") or "").strip(),
        )
        current = quota_by_subscription.get(key)
        if current is None:
            current = {
                "tenant_id": (row.get("tenant_id") or "").strip(),
                "regions": set(),
                "max_utilization_pct": 0.0,
                "high_quota_count": 0,
            }
            quota_by_subscription[key] = current

        location = (row.get("location") or "").strip()
        if location:
            current["regions"].add(location)

        pct = _safe_float(row.get("utilization_pct"))
        if pct > current["max_utilization_pct"]:
            current["max_utilization_pct"] = pct
        if pct >= 80.0:
            current["high_quota_count"] += 1

    subscriptions_quotas_rows = []
    for (sub_name, sub_id), stats in sorted(quota_by_subscription.items(), key=lambda item: item[0]):
        subscriptions_quotas_rows.append([
            sub_name,
            sub_id,
            stats["tenant_id"],
            len(stats["regions"]),
            f"{_safe_float(stats['max_utilization_pct']):.1f}",
            stats["high_quota_count"],
        ])

    entra_secrets_expired = [
        row for row in entra_rows
        if (row.get("status") or "").strip().lower() == "expired"
    ]
    entra_secrets_expiring_90d = [
        row for row in entra_rows
        if (row.get("status") or "").strip().lower() == "expiring"
        and 0 <= _safe_int(row.get("days_remaining")) <= 90
    ]
    entra_secrets_expiring_180d = [
        row for row in entra_rows
        if (row.get("status") or "").strip().lower() in {"expiring", "valid"}
        and 91 <= _safe_int(row.get("days_remaining")) <= 180
    ]
    entra_secrets_noexpiry = [
        row for row in entra_rows
        if not (row.get("end_utc") or "").strip()
    ]

    def _entra_table_rows(rows: list[dict[str, str]], limit: int = 100) -> list[list[str]]:
        data = []
        for row in rows[:limit]:
            data.append([
                row.get("tenant_id", ""),
                row.get("app_display_name", ""),
                row.get("credential_type", ""),
                row.get("credential_display_name", ""),
                row.get("days_remaining", ""),
                row.get("end_utc", ""),
            ])
        return data

    payload = {
        "scalars": {
            "CUSTOMER": customer,
            "DATE": date.today().strftime("%d %B %Y"),
            "SCAN_DATE": str(manifest.get("generated_at", date.today().isoformat())),
            "TENANT": tenant_scalar,
            "SCAN_OPERATOR": os.environ.get("USERNAME") or os.environ.get("USER") or "Unknown",
        },
        "tables": {
            "SUBSCRIPTIONS_IN_SCOPE": {
                "headers": ["Subscription Name", "Subscription ID", "Management group"],
                "rows": [
                    [row.get("subscription_name", ""), row.get("subscription_id", ""), row.get("management_group", "")]
                    for row in in_scope
                ] or [["(none)", "", ""]],
            },
            "SUBSCRIPTIONS_OUT_OF_SCOPE": {
                "headers": ["Subscription Name", "Subscription ID", "Reason / Status"],
                "rows": [
                    [row.get("subscription_name", ""), row.get("subscription_id", ""), row.get("subscription_status", "")]
                    for row in out_scope
                ] or [["(none)", "", ""]],
            },
            "SUBSCRIPTIONS_QUOTAS": {
                "headers": [
                    "Subscription Name",
                    "Subscription ID",
                    "Tenant ID",
                    "Regions with quota data",
                    "Max utilization %",
                    "Quota items >=80%",
                ],
                "rows": subscriptions_quotas_rows or [["(no data)", "", "", "", "", ""]],
            },
            "APPSERVICE": {
                "headers": [
                    "ASP Name",
                    "SKU & Instances",
                    "Number of applications",
                    "Cost signal",
                    "CPU p95%",
                    "Memory p95%",
                ],
                "rows": [
                    [
                        row.get("plan_name") or row.get("name", ""),
                        _format_asp_sku_instances(row),
                        row.get("apps_count", ""),
                        row.get("cost_signal_category", ""),
                        row.get("cpu_p95_pct_window", ""),
                        row.get("memory_p95_pct_window", ""),
                    ]
                    for row in appservice_rows[:50]
                ] or [["(no data)", "", "", "", "", ""]],
            },
            "KEYVAULT": {
                "headers": ["Vault", "Resource Group", "SKU", "RBAC", "Private EP", "Quality"],
                "rows": [
                    [
                        row.get("key_vault_name", ""),
                        row.get("resource_group", ""),
                        row.get("sku_name", ""),
                        row.get("qc_rbac_authorization", row.get("enable_rbac_authorization", "")),
                        row.get("qc_private_endpoint_configured", row.get("private_endpoint_count", "")),
                        row.get("quality_score", ""),
                    ]
                    for row in keyvault_rows[:80]
                ] or [["(no data)", "", "", "", "", ""]],
            },
            "STORAGE_ACCOUNTS": {
                "headers": ["Account", "RG", "SKU", "TLS", "HTTPS only", "Public Access", "HNS", "Quality"],
                "rows": [
                    [
                        row.get("storage_account_name", ""),
                        row.get("resource_group", ""),
                        row.get("sku_name", ""),
                        row.get("minimum_tls_version", ""),
                        row.get("https_only", ""),
                        row.get("public_network_access_enabled", ""),
                        row.get("is_hns_enabled", ""),
                        row.get("quality_score", ""),
                    ]
                    for row in storage_rows[:80]
                ] or [["(no data)", "", "", "", "", "", "", ""]],
            },
            "SQL": {
                "headers": ["Server", "Database", "Resource Group", "Tier", "CPU p95 %", "Signal"],
                "rows": [
                    [
                        row.get("server_name", ""),
                        row.get("database_name", ""),
                        row.get("resource_group", ""),
                        row.get("sku_tier", ""),
                        row.get("cpu_p95_pct", ""),
                        row.get("advisory_category", ""),
                    ]
                    for row in sql_rows[:80]
                ] or [["(no data)", "", "", "", "", ""]],
            },
            "ENTRA_SECRETS_EXPIRED": {
                "headers": ["Tenant", "Application", "Credential Type", "Credential Name", "Days Remaining", "End (UTC)"],
                "rows": _entra_table_rows(entra_secrets_expired) or [["(none)", "", "", "", "", ""]],
            },
            "ENTRA_SECRETS_EXPIRING_90D": {
                "headers": ["Tenant", "Application", "Credential Type", "Credential Name", "Days Remaining", "End (UTC)"],
                "rows": _entra_table_rows(entra_secrets_expiring_90d) or [["(none)", "", "", "", "", ""]],
            },
            "ENTRA_SECRETS_EXPIRING_180D": {
                "headers": ["Tenant", "Application", "Credential Type", "Credential Name", "Days Remaining", "End (UTC)"],
                "rows": _entra_table_rows(entra_secrets_expiring_180d) or [["(none)", "", "", "", "", ""]],
            },
            "ENTRA_SECRETS_NOEXPIRY": {
                "headers": ["Tenant", "Application", "Credential Type", "Credential Name", "Days Remaining", "End (UTC)"],
                "rows": _entra_table_rows(entra_secrets_noexpiry) or [["(none)", "", "", "", "", ""]],
            },
            "APP_INSIGHTS": {
                "headers": ["Workspace", "Resource Group", "SKU", "Daily Cap (GB)", "Cap Exceeded", "Retention (days)", "Avg 7d Ingest (GB)", "Created"],
                "rows": [
                    [
                        row.get("workspace_name", ""),
                        row.get("resource_group", ""),
                        row.get("sku_name", ""),
                        row.get("daily_quota_gb", ""),
                        row.get("daily_quota_exceeded", ""),
                        row.get("retention_days", ""),
                        row.get("ingestion_gb_last_7d_avg", ""),
                        (row.get("created_date") or "")[:10],
                    ]
                    for row in law_rows[:80]
                ] or [["(no data)", "", "", "", "", "", "", ""]],
            },
        },
        "recommendations": {
            "KEYVAULT": (
                f"Key Vaults with quality_score < 60: {len(keyvault_low_quality)} out of {len(keyvault_rows)}.\n"
                "Prioritize purge protection, RBAC authorization, and private endpoint coverage."
            ),
            "APPSERVICE": (
                f"App Service plans flagged for review: {len(appservice_reco)} out of {len(appservice_rows)}.\n"
                "Cost signal status legend:\n"
                "- reviewed: baseline review recommended (check configuration, utilization, and sizing).\n"
                "- stopped: hosted apps are stopped; evaluate shutdown/decommission opportunities.\n"
                "- scale-up-review: sustained pressure observed; review for potential scale-up.\n"
                "- scale-in-candidate: low sustained utilization; review for potential scale-in/right-sizing.\n"
                "- oversized-premium: premium tier appears overprovisioned for observed demand; evaluate SKU downgrade.\n"
                "- orphaned: plan has no hosted apps; remove if no longer required."
            ),
            "STORAGE_ACCOUNTS": (
                f"Storage accounts with quality_score < 60: {len(storage_low_quality)} out of {len(storage_rows)}.\n"
                "Prioritize HTTPS-only enforcement, firewall default deny, and soft-delete/versioning controls."
            ),
            "SQL": (
                f"SQL databases with pressure signals: {len(sql_pressure)} out of {len(sql_rows)}.\n"
                "Review scale-up, storage pressure, deadlocks, and ensure TDE and backup retention policies are compliant."
            ),
            "ENTRA_SECRETS": (
                f"Entra credential posture: expired={len(entra_secrets_expired)}, "
                f"expiring<=90d={len(entra_secrets_expiring_90d)}, "
                f"expiring 91-180d={len(entra_secrets_expiring_180d)}, "
                f"no-expiry={len(entra_secrets_noexpiry)}.\n"
                "Immediate actions:\n"
                "- Rotate all expired secrets/certificates and validate app health after rotation.\n"
                "- Schedule rolling rotation for <=90 day credentials and move to certificate-based auth where possible.\n"
                "- Review 91-180 day credentials for planned replacement windows.\n"
                "- Investigate no-expiry credentials and enforce finite expiry plus owner/rotation policy."
            ),
            "APP_INSIGHTS": _build_law_recommendation(law_rows),
        },
        "dashboard_tables": _build_dashboard_tables(manifest, export_dir),
    }

    return payload


def _safe_float(value: str | None) -> float:
    if value is None:
        return 0.0
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def _build_dashboard_tables(manifest: dict, export_dir: Path) -> dict:
    """Create generic table specs for each dashboard summary file in manifest."""
    tables: dict = {}
    for dashboard in manifest.get("dashboards", []):
        dash_id = str(dashboard.get("id", ""))
        summary_name = None
        for file_spec in dashboard.get("files", []):
            if file_spec.get("type") == "summary":
                summary_name = file_spec.get("filename")
                break
        if not summary_name:
            continue

        rows = _read_csv(export_dir / summary_name)
        if not rows:
            continue

        headers = list(rows[0].keys())[:8]
        data_rows = [[row.get(h, "") for h in headers] for row in rows[:60]]
        key = _normalize_token(dash_id)
        tables[key] = {"headers": headers, "rows": data_rows}

    fallback_sources = {
        "ACR": "acr_summary",
        "APIM": "apim_summary",
        "APPCONFIG": "app_config_summary",
        "APPINSIGHTS": "app_insights_summary",
        "LOGANALYTICSWORKSPACES": "loganalyticsworkspace_summary",
        "APPSERVICEPLANS": "app_service_plans_summary",
        "COGNITIVE": "cognitive_summary",
        "CONTAINERAPPS": "container_apps_summary",
        "COSMOSDB": "cosmosdb_summary",
        "EVENTHUBSB": "eventhub_summary",
        "KEYVAULTS": "keyvaults_summary",
        "LOGICDF": "logic_df_summary",
        "POSTGRESQL": "postgresql_summary",
        "SQLDATABASES": "sql_summary",
        "SQL": "sql_summary",
        "VIRTUALMACHINES": "virtual_machines_summary",
        "EVENTHUBS": "eventhub_summary",
        "AZURECOSTS": "daily_costs",
        "STORAGEACCOUNTS": "storage_accounts_summary",
    }

    for key, fragment in fallback_sources.items():
        if key in tables:
            continue
        csv_path = _find_latest_csv_by_fragment(export_dir, fragment)
        if not csv_path:
            continue
        rows = _read_csv(csv_path)
        if not rows:
            continue
        headers = list(rows[0].keys())[:8]
        data_rows = [[row.get(h, "") for h in headers] for row in rows[:60]]
        tables[key] = {"headers": headers, "rows": data_rows}

    return tables


def _table_spec_for_anchor(anchor_key: str, payload: dict):
    """Resolve a table payload for a template TABLE_* anchor."""
    direct = payload["tables"].get(anchor_key)
    if direct:
        return direct

    normalized = _normalize_token(anchor_key)

    aliases = {
        "ACR": "ACR",
        "APIM": "APIM",
        "APPCONFIG": "APPCONFIG",
        "APPINSIGHTS": "APPINSIGHTS",
        "APPSERVICE": "APPSERVICE",
        "STORAGE": "STORAGE_ACCOUNTS",
        "STORAGEACCOUNT": "STORAGE_ACCOUNTS",
        "STORAGEACCOUNTS": "STORAGE_ACCOUNTS",
        "KEYVAULT": "KEYVAULT",
        "KEYVAULTS": "KEYVAULT",
        "APPSERVICEPLANS": "APPSERVICE",
        "APPSERVICEPLAN": "APPSERVICE",
        "COGNITIVE": "COGNITIVE",
        "COSMOSDB": "COSMOSDB",
        "EVENTHUBSB": "EVENTHUBS",
        "VM": "VIRTUALMACHINES",
        "VMS": "VIRTUALMACHINES",
        "LOGICDF": "LOGICDF",
        "LOGICDATAFACTORY": "LOGICDF",
        "OVERALLPOSTURE": "OVERALL_POSTURE",
        "AGREEDACTIONS": "AGREED_ACTIONS",
        "SQL": "SQL",
        "SQLDATABASE": "SQL",
        "SQLDATABASES": "SQL",
        "CONTAINERAPPSAKS": "CONTAINERAPPS",
        "EVENTHUBSB": "EVENTHUBS",
    }
    target = aliases.get(normalized)
    if target:
        if payload["tables"].get(target):
            return payload["tables"][target]
        if payload.get("dashboard_tables", {}).get(target):
            return payload["dashboard_tables"][target]

    for key, value in payload["tables"].items():
        if _normalize_token(key) == normalized:
            return value

    dashboard_alias = {
        "KEYVAULT": "KEYVAULTS",
        "KEYVAULTS": "KEYVAULTS",
        "VIRTUALMACHINES": "VIRTUALMACHINES",
        "POSTGRESQL": "POSTGRESQL",
        "SQL": "SQL",
        "SQLDATABASE": "SQLDATABASES",
        "SQLDATABASES": "SQLDATABASES",
        "CONTAINERAPPS": "CONTAINERAPPS",
        "EVENTHUBS": "EVENTHUBS",
        "AZURECOSTS": "AZURECOSTS",
    }
    dashboard_key = dashboard_alias.get(normalized, normalized)
    return payload.get("dashboard_tables", {}).get(dashboard_key)


def _recommendation_for_anchor(anchor_key: str, payload: dict) -> str | None:
    direct = payload["recommendations"].get(anchor_key)
    if direct:
        return direct
    normalized = _normalize_token(anchor_key)
    reco_aliases = {
        "ACR": "ACR",
        "APIM": "APIM",
        "APPCONFIG": "APP_CONFIG",
        "APPINSIGHTS": "APP_INSIGHTS",
        "COGNITIVE": "COGNITIVE",
        "CONTAINERAPPSAKS": "CONTAINER_APPS_AKS",
        "COSMOSDB": "COSMOSDB",
        "EVENTHUBSB": "EVENTHUB_SB",
        "EVENTHUBSERVICEBUS": "EVENTHUB_SB",
        "KEYVAULT": "KEYVAULT",
        "KEYVAULTS": "KEYVAULT",
        "LOGICDF": "LOGIC_DF",
        "APPSERVICEPLAN": "APPSERVICE",
        "APPSERVICE": "APPSERVICE",
        "STORAGE": "STORAGE_ACCOUNTS",
        "STORAGEACCOUNT": "STORAGE_ACCOUNTS",
        "STORAGEACCOUNTS": "STORAGE_ACCOUNTS",
        "VM": "VM",
        "VMS": "VM",
        "SQL": "SQL",
        "SQLDATABASE": "SQL",
        "SQLDATABASES": "SQL",
    }
    alias_key = reco_aliases.get(normalized)
    if alias_key and payload["recommendations"].get(alias_key):
        return payload["recommendations"][alias_key]

    for key, value in payload["recommendations"].items():
        if _normalize_token(key) == normalized:
            return value
    return None


def _fallback_table_spec(anchor_key: str) -> dict:
    """Return a default table payload when a template anchor has no data mapping."""
    return {
        "headers": ["Status", "Details"],
        "rows": [["No data mapping", f"No extractor output is currently mapped for TABLE_{anchor_key}."]],
    }


def _fallback_recommendation(anchor_key: str) -> str:
    """Return a default recommendation for unmapped recommendation anchors."""
    return (
        f"No automated recommendation is currently mapped for RECOMMENDATIONS_{anchor_key}.\n"
        "Please review this section manually based on customer-specific context."
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Fill Quality Controls DOCX from latest managed-service exports.")
    parser.add_argument("-c", "--customer", required=True, help="Customer code, e.g. CUST")
    parser.add_argument("--template", default=None, help="Path to QC template DOCX (overrides version selection)")
    parser.add_argument(
        "--template-version",
        default=None,
        help=(
            "Template version suffix after _v (for example: 1.1 or SYNH). "
            "When omitted, the highest numeric version in templates/ is selected automatically."
        ),
    )
    parser.add_argument("--output", default=None, help="Output DOCX path")
    parser.add_argument("--reports-dir", default=str(DEFAULT_REPORTS_DIR), help="Root reports directory")
    # -- Storage download (preferred source of input) -------------------------
    parser.add_argument(
        "--storage-connection-string",
        default=None,
        help="SAS URL or connection string to download report files from Azure Blob Storage "
             "before filling the template. When provided, the wrapper is NOT executed.",
    )
    parser.add_argument(
        "--storage-container",
        default="managed-service-reports",
        help="Blob container to download from (default: managed-service-reports).",
    )
    parser.add_argument(
        "--storage-prefix",
        default=None,
        help="Blob path prefix inside the container (default: customer code).",
    )
    # -- Legacy / explicit wrapper execution ----------------------------------
    parser.add_argument(
        "--run-wrapper",
        action="store_true",
        help="Explicitly run managedServiceWrapper.py to refresh extracts before filling the template.",
    )
    parser.add_argument("--skip-login", action="store_true", help="Pass --skip-login to managedServiceWrapper")
    parser.add_argument("--from", dest="date_from", default=None, help="Pass-through start date for wrapper")
    parser.add_argument("--to", dest="date_to", default=None, help="Pass-through end date for wrapper")
    parser.add_argument("--lookback", default=None, help="Pass-through lookback for wrapper")
    parser.add_argument("--sp-client-id", default=None, help="Service principal client id")
    parser.add_argument("--sp-client-secret", default=os.environ.get("AZURE_SP_CLIENT_SECRET"), help="Service principal client secret")
    parser.add_argument("--sp-certificate", default=None, help="Service principal certificate path")
    # -- Backward compat (no-op; wrapper is now opt-in via --run-wrapper) -----
    parser.add_argument("--skip-wrapper", action="store_true", help=argparse.SUPPRESS)
    return parser.parse_args()


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------

def main():
    args = parse_args()

    customer = args.customer.upper()
    template_path = _resolve_template_path(args.template, args.template_version)
    reports_dir = Path(args.reports_dir)
    output_path = Path(args.output) if args.output else Path(f"{date.today().isoformat()}_{customer}_QualityControls.docx")

    if not template_path.exists():
        raise SystemExit(f"Template not found: {template_path}")

    if args.run_wrapper:
        # Explicit opt-in: run all extractors fresh.
        _run_wrapper(
            customer=customer,
            reports_dir=reports_dir,
            skip_login=args.skip_login,
            date_from=args.date_from,
            date_to=args.date_to,
            lookback=args.lookback,
            sp_client_id=args.sp_client_id,
            sp_client_secret=args.sp_client_secret,
            sp_certificate=args.sp_certificate,
        )
    elif args.storage_connection_string:
        # Download manifest + report files from blob storage.
        storage_prefix = args.storage_prefix or customer
        print(f"Downloading reports from blob storage (prefix: {storage_prefix})...")
        _download_reports_from_blob(
            customer=customer,
            reports_dir=reports_dir,
            connection_input=args.storage_connection_string,
            container_name=args.storage_container,
            blob_prefix=storage_prefix,
        )
    # else: use whatever is already present locally in reports_dir.

    export_dir = _find_latest_export_dir(customer, reports_dir)
    manifest_path = export_dir / "manifest.json"
    if not manifest_path.exists():
        raise SystemExit(f"manifest.json not found in export folder: {export_dir}")

    manifest = _read_json(manifest_path)
    data = _build_payload_from_exports(customer, export_dir, manifest)

    print(f"Using export folder: {export_dir}")
    print(f"Using manifest: {manifest_path}")
    print(f"Using template: {template_path}")

    doc = Document(str(template_path))

    # 1. Scalars (cover page, scan metadata, change-history date cell, etc.)
    for key, value in data["scalars"].items():
        n = replace_scalar(doc, key, value)
        print(f"  scalar {key:25s} -> {n} replacement(s)")

    # 2. Tables (driven by template anchors + manifest-backed data)
    table_anchors = _iter_anchor_keys(doc, "TABLE")
    for anchor in table_anchors:
        short_key = anchor[len("TABLE_"):]
        spec = _table_spec_for_anchor(short_key, data)
        if not spec:
            spec = _fallback_table_spec(short_key)
            ok = insert_table_at(doc, anchor, spec["headers"], spec["rows"])
            print(f"  table  {anchor:30s} -> {'OK (fallback)' if ok else 'ANCHOR NOT FOUND'}")
            continue

        ok = insert_table_at(doc, anchor, spec["headers"], spec["rows"])
        print(f"  table  {anchor:30s} -> {'OK' if ok else 'ANCHOR NOT FOUND'}")

    # 3. Recommendations (driven by template anchors)
    reco_anchors = _iter_anchor_keys(doc, "RECOMMENDATIONS")
    for anchor in reco_anchors:
        short_key = anchor[len("RECOMMENDATIONS_"):]
        text = _recommendation_for_anchor(short_key, data)
        if not text:
            text = _fallback_recommendation(short_key)
            ok = replace_recommendation(doc, anchor, text)
            print(f"  reco   {anchor:30s} -> {'OK (fallback)' if ok else 'ANCHOR NOT FOUND'}")
            continue

        ok = replace_recommendation(doc, anchor, text)
        print(f"  reco   {anchor:30s} -> {'OK' if ok else 'ANCHOR NOT FOUND'}")

    doc.save(str(output_path))
    print(f"\nWrote {output_path}")
    print("Open in Word and press F9 (or right-click the TOC -> Update Field) to refresh the table of contents.")


if __name__ == "__main__":
    main()
